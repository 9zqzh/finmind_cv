r"""一次性脚本：把 resources/ 下的 PDF/docx 转成 Markdown 并分类放入知识库目录。

用法（在 backend 目录下）：
    .venv\Scripts\python _convert_resources.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent.parent.parent  # 项目根目录
RESOURCES = ROOT / "resources"
KNOWLEDGE = Path(__file__).resolve().parent / "data" / "knowledge"

# 关键词检索依赖原词命中：为文档补充用户常用的同义问法，提高命中率
ALIASES: dict[str, list[str]] = {
    "附件2：清远校区外来人员临时进校审批表": [
        "临时入校流程",
        "外来人员入校申请",
        "访客进校审批",
        "车辆进校申请",
    ],
}


def clean_text(text: str) -> str:
    """压缩多余空行、去掉行尾空白。"""
    lines = [line.rstrip() for line in text.splitlines()]
    out: list[str] = []
    blank = 0
    for line in lines:
        if line.strip():
            out.append(line)
            blank = 0
        else:
            blank += 1
            if blank <= 1:
                out.append("")
    return "\n".join(out).strip()


_SECTION_START = re.compile(r"^[一二三四五六七八九十\d（(【\[]")
_SENTENCE_END = "。；：！？）】”"

# 制度类文档结构化：章/条标题提升为 Markdown 标题与独立段落，让检索按条款分块
_PAGE_NUM_RE = re.compile(r"^-\s*\d+\s*-", re.MULTILINE)
_CHAPTER_RE = re.compile(r"(第[一二三四五六七八九十百]+章\s*[^\n第]{1,20})")
_ARTICLE_RE = re.compile(r"(第[一二三四五六七八九十百]+条)")


def structure_policy_text(text: str) -> str:
    """为制度类文本补充结构：去页码、章提升为标题、条款独立成段。"""
    text = _PAGE_NUM_RE.sub("", text)
    text = _CHAPTER_RE.sub(r"\n## \1\n", text)
    text = _ARTICLE_RE.sub(r"\n\n\1", text)
    return text


def join_wrapped_lines(text: str) -> str:
    """把 PDF 提取出的句中断行拼接回完整句子（启发式）。"""
    lines = [line.rstrip() for line in text.splitlines()]
    merged: list[str] = []
    for line in lines:
        if not line.strip():
            merged.append(line)
            continue
        if merged and merged[-1].strip():
            prev = merged[-1]
            prev_end_ok = prev and prev[-1] not in _SENTENCE_END
            next_start_ok = not _SECTION_START.match(line)
            if prev_end_ok and next_start_ok:
                merged[-1] = prev + line
                continue
        merged.append(line)
    return "\n".join(merged)


def pdf_to_markdown(pdf_path: Path, title: str) -> str:
    reader = PdfReader(str(pdf_path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        page_text = page.extract_text() or ""
        page_text = join_wrapped_lines(page_text).strip()
        if page_text:
            parts.append(f"<!-- 第 {i} 页 -->\n\n{page_text}")
    body = clean_text("\n\n".join(parts))
    # 制度细则类文档（含“第X章/第X条”）做结构化处理，便于按条款分块检索
    if _CHAPTER_RE.search(body) or _ARTICLE_RE.search(body):
        body = structure_policy_text(body)
        body = clean_text(body)
    aliases = ALIASES.get(title)
    if aliases:
        body += "\n\n相关关键词：" + "、".join(aliases)
    header = (
        f"# {title}\n\n"
        f"> 来源文件：resources/{pdf_path.relative_to(RESOURCES).as_posix()}\n"
        f"> 页数：{len(reader.pages)}\n"
    )
    return header + "\n" + body + "\n"


def _dedupe_row(cells: list[str]) -> list[str]:
    """合并单元格会被 python-docx 重复返回：去连续重复，去空值。"""
    out: list[str] = []
    for cell in cells:
        cell = cell.strip().replace("\n", " ")
        if not cell or (out and out[-1] == cell):
            continue
        out.append(cell)
    return out


def docx_to_markdown(docx_path: Path, title: str) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        # 表单类表格转成“字段：值”列表，比原始管道符表格更易读
        rows: list[str] = []
        for row in table.rows:
            cells = _dedupe_row([cell.text for cell in row.cells])
            if cells:
                rows.append("- " + "：".join(cells))
        if rows:
            parts.append("\n".join(rows))
    body = clean_text("\n\n".join(parts))
    aliases = ALIASES.get(title)
    if aliases:
        body += "\n\n相关关键词：" + "、".join(aliases)
    header = (
        f"# {title}\n\n"
        f"> 来源文件：resources/{docx_path.relative_to(RESOURCES).as_posix()}\n"
    )
    return header + "\n" + body + "\n"


def safe_name(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', "", name).strip()


def main() -> None:
    plan_dir = KNOWLEDGE / "培养方案"
    policy_dir = KNOWLEDGE / "制度流程"
    plan_dir.mkdir(parents=True, exist_ok=True)
    policy_dir.mkdir(parents=True, exist_ok=True)

    stats: list[tuple[str, str, int]] = []

    # 1. 人才培养方案 PDF -> data/knowledge/培养方案/
    for pdf in sorted(p for p in (RESOURCES / "人才培养方案").glob("*.pdf") if not p.name.startswith("~$")):
        title = pdf.stem
        md = pdf_to_markdown(pdf, title)
        out = plan_dir / f"{safe_name(title)}.md"
        out.write_text(md, encoding="utf-8")
        stats.append((str(out.relative_to(ROOT)), title, len(md)))

    # 2. 顶层制度类 PDF -> data/knowledge/制度流程/
    for pdf in sorted(p for p in RESOURCES.glob("*.pdf") if not p.name.startswith("~$")):
        title = pdf.stem
        md = pdf_to_markdown(pdf, title)
        out = policy_dir / f"{safe_name(title)}.md"
        out.write_text(md, encoding="utf-8")
        stats.append((str(out.relative_to(ROOT)), title, len(md)))

    # 3. docx 表格 -> data/knowledge/制度流程/（跳过 Word 打开时产生的 ~$ 锁文件）
    for docx in sorted(p for p in RESOURCES.glob("*.docx") if not p.name.startswith("~$")):
        title = docx.stem
        md = docx_to_markdown(docx, title)
        out = policy_dir / f"{safe_name(title)}.md"
        out.write_text(md, encoding="utf-8")
        stats.append((str(out.relative_to(ROOT)), title, len(md)))

    for path, title, size in stats:
        print(f"[OK] {title} -> {path}（{size} 字符）")


if __name__ == "__main__":
    main()

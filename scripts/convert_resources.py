r"""一次性脚本：把 resources/ 下的 PDF/docx/doc/xlsx 转成 Markdown 并分类放入知识库目录。

用法（在项目根目录下）：
    backend\.venv\Scripts\python scripts\convert_resources.py

支持的格式：
- .pdf  → pypdf 提取文本
- .docx → python-docx 提取段落与表格
- .doc  → pywin32 COM 自动化（需安装 Word）
- .xlsx → openpyxl 提取工作表内容

已存在的 .md 文件默认跳过，不会重复转换。
"""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent.parent  # 项目根目录
RESOURCES = ROOT / "resources"
KNOWLEDGE = ROOT / "backend" / "data" / "knowledge"

# 关键词检索依赖原词命中：为文档补充用户常用的同义问法，提高命中率
ALIASES: dict[str, list[str]] = {
    "附件2：清远校区外来人员临时进校审批表": [
        "临时入校流程",
        "外来人员入校申请",
        "访客进校审批",
        "车辆进校申请",
    ],
    "学生申请课程重修流程": [
        "课程重修",
        "重修怎么申请",
        "重修流程",
        "怎么重修",
    ],
    "学生重修报名指南": [
        "重修报名",
        "重修怎么报名",
        "重修报名流程",
    ],
    "广东金融学院学生评教操作手册（简化版）": [
        "评教",
        "学生评教",
        "评教怎么操作",
        "评教流程",
    ],
    "跟班重修申请表": [
        "跟班重修",
        "重修申请表",
        "重修表怎么填",
    ],
    "附：广东金融学院2026-2027学年第一学期校历": [
        "校历",
        "学期校历",
        "开学时间",
        "放假时间",
        "考试周",
    ],
    "附件2：广东金融学院学分制综合教务管理系统学生网上操作手册（选课指南）": [
        "选课指南",
        "选课操作",
        "怎么选课",
        "教务系统操作",
        "网上选课",
    ],
    "关于做好2025-2026学年第二学期课程重修报名工作的通知": [
        "重修报名通知",
        "重修报名时间",
        "课程重修通知",
    ],
    "关于组织清远校区2024、2025级学生进行2025-2026学年第二学期选课的通知": [
        "选课通知",
        "清远校区选课",
        "第二学期选课",
    ],
    "附件1：选课时间安排": [
        "选课时间",
        "选课安排",
        "什么时候选课",
    ],
}


def clean_text(text: str) -> str:
    """压缩多余空行、去掉行尾空白。"""
    lines = [line.rstrip() for line in text.splitlines()]
    out: list[str] = []
    blank = 0
    for line in lines:
        if line.strip():
            out.append(line)
            blank = 0
        else:
            blank += 1
            if blank <= 1:
                out.append("")
    return "\n".join(out).strip()


_SECTION_START = re.compile(r"^[一二三四五六七八九十\d（(【\[]")
_SENTENCE_END = "。；：！？）】”"

# 制度类文档结构化：章/条标题提升为 Markdown 标题与独立段落，让检索按条款分块
_PAGE_NUM_RE = re.compile(r"^-\s*\d+\s*-", re.MULTILINE)
_CHAPTER_RE = re.compile(r"(第[一二三四五六七八九十百]+章\s*[^\n第]{1,20})")
_ARTICLE_RE = re.compile(r"(第[一二三四五六七八九十百]+条)")


def structure_policy_text(text: str) -> str:
    """为制度类文本补充结构：去页码、章提升为标题、条款独立成段。"""
    text = _PAGE_NUM_RE.sub("", text)
    text = _CHAPTER_RE.sub(r"\n## \1\n", text)
    text = _ARTICLE_RE.sub(r"\n\n\1", text)
    return text


def join_wrapped_lines(text: str) -> str:
    """把 PDF 提取出的句中断行拼接回完整句子（启发式）。"""
    lines = [line.rstrip() for line in text.splitlines()]
    merged: list[str] = []
    for line in lines:
        if not line.strip():
            merged.append(line)
            continue
        if merged and merged[-1].strip():
            prev = merged[-1]
            prev_end_ok = prev and prev[-1] not in _SENTENCE_END
            next_start_ok = not _SECTION_START.match(line)
            if prev_end_ok and next_start_ok:
                merged[-1] = prev + line
                continue
        merged.append(line)
    return "\n".join(merged)


def pdf_to_markdown(pdf_path: Path, title: str) -> str:
    reader = PdfReader(str(pdf_path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        page_text = page.extract_text() or ""
        page_text = join_wrapped_lines(page_text).strip()
        if page_text:
            parts.append(f"<!-- 第 {i} 页 -->\n\n{page_text}")
    body = clean_text("\n\n".join(parts))
    # 制度细则类文档（含“第X章/第X条”）做结构化处理，便于按条款分块检索
    if _CHAPTER_RE.search(body) or _ARTICLE_RE.search(body):
        body = structure_policy_text(body)
        body = clean_text(body)
    aliases = ALIASES.get(title)
    if aliases:
        body += "\n\n相关关键词：" + "、".join(aliases)
    header = (
        f"# {title}\n\n"
        f"> 来源文件：resources/{pdf_path.relative_to(RESOURCES).as_posix()}\n"
        f"> 页数：{len(reader.pages)}\n"
    )
    return header + "\n" + body + "\n"


def _dedupe_row(cells: list[str]) -> list[str]:
    """合并单元格会被 python-docx 重复返回：去连续重复，去空值。"""
    out: list[str] = []
    for cell in cells:
        cell = cell.strip().replace("\n", " ")
        if not cell or (out and out[-1] == cell):
            continue
        out.append(cell)
    return out


def docx_to_markdown(docx_path: Path, title: str) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        # 表单类表格转成“字段：值”列表，比原始管道符表格更易读
        rows: list[str] = []
        for row in table.rows:
            cells = _dedupe_row([cell.text for cell in row.cells])
            if cells:
                rows.append("- " + "：".join(cells))
        if rows:
            parts.append("\n".join(rows))
    body = clean_text("\n\n".join(parts))
    aliases = ALIASES.get(title)
    if aliases:
        body += "\n\n相关关键词：" + "、".join(aliases)
    header = (
        f"# {title}\n\n"
        f"> 来源文件：resources/{docx_path.relative_to(RESOURCES).as_posix()}\n"
    )
    return header + "\n" + body + "\n"


def safe_name(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', "", name).strip()


def doc_to_markdown(doc_path: Path, title: str) -> str:
    """把旧版 .doc 文件通过 Word COM 自动化转为文本再包装成 Markdown。

    需要本机安装 Microsoft Word；若 Word 不可用则返回 None。
    """
    try:
        import win32com.client
    except ImportError:
        print(f"[SKIP] {title}：pywin32 未安装，无法处理 .doc 文件")
        return None

    word = None
    doc_obj = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_obj = word.Documents.Open(str(doc_path.resolve()))
        # 导出为纯文本（wdFormatText = 2），Word 默认用系统编码（中文 Windows 为 GBK）
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = tmp.name
        doc_obj.SaveAs(tmp_path, FileFormat=2, Encoding=65001)
        # 先关闭文档释放文件锁，再读取临时文件
        doc_obj.Close(False)
        doc_obj = None
        raw = Path(tmp_path).read_bytes()
        try:
            Path(tmp_path).unlink(missing_ok=True)
        except OSError:
            pass  # 文件锁未完全释放时忽略
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("gbk", errors="replace")
    except Exception as exc:
        print(f"[SKIP] {title}：Word COM 转换失败：{exc}")
        return None
    finally:
        if doc_obj:
            try:
                doc_obj.Close(False)
            except Exception:
                pass
        if word:
            try:
                word.Quit()
            except Exception:
                pass

    body = clean_text(join_wrapped_lines(text))
    aliases = ALIASES.get(title)
    if aliases:
        body += "\n\n相关关键词：" + "、".join(aliases)
    header = (
        f"# {title}\n\n"
        f"> 来源文件：resources/{doc_path.relative_to(RESOURCES).as_posix()}\n"
    )
    return header + "\n" + body + "\n"


def xlsx_to_markdown(xlsx_path: Path, title: str) -> str:
    """把 xlsx 各工作表内容提取为 Markdown（表格 + 文本混合）。"""
    wb = load_workbook(str(xlsx_path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if any(cells):
                rows_data.append(cells)
        if not rows_data:
            continue
        parts.append(f"## {sheet_name}")
        # 把表格转为 Markdown 表格格式
        max_cols = max(len(r) for r in rows_data)
        # 补齐列数
        for r in rows_data:
            while len(r) < max_cols:
                r.append("")
        # 第一行作为表头
        header_row = "| " + " | ".join(rows_data[0]) + " |"
        separator = "| " + " | ".join(["---"] * max_cols) + " |"
        data_rows = []
        for r in rows_data[1:]:
            data_rows.append("| " + " | ".join(r) + " |")
        parts.append("\n".join([header_row, separator] + data_rows))
    wb.close()

    body = clean_text("\n\n".join(parts))
    aliases = ALIASES.get(title)
    if aliases:
        body += "\n\n相关关键词：" + "、".join(aliases)
    header = (
        f"# {title}\n\n"
        f"> 来源文件：resources/{xlsx_path.relative_to(RESOURCES).as_posix()}\n"
    )
    return header + "\n" + body + "\n"


def _already_converted(out_path: Path, src_path: Path) -> bool:
    """如果输出 .md 已存在且比源文件新，跳过。"""
    if not out_path.exists():
        return False
    return out_path.stat().st_mtime >= src_path.stat().st_mtime


def main() -> None:
    # 知识库子目录映射：resources 子目录 → knowledge 子目录
    SUBDIR_MAP = {
        "人才培养方案": "培养方案",
        "办事流程": "办事流程",
        "教学安排": "教学安排",
        "规章制度": "制度流程",
    }

    # 确保所有目标子目录存在
    for target in SUBDIR_MAP.values():
        (KNOWLEDGE / target).mkdir(parents=True, exist_ok=True)
    policy_dir = KNOWLEDGE / "制度流程"

    stats: list[tuple[str, str, int]] = []
    skipped: list[str] = []

    def _convert_one(src: Path, target_subdir: str) -> None:
        """根据扩展名转换单个文件，写入 target_subdir 对应的知识库目录。"""
        title = src.stem
        out_dir = KNOWLEDGE / target_subdir
        out_name = safe_name(title) + ".md"
        out_path = out_dir / out_name

        if _already_converted(out_path, src):
            skipped.append(title)
            return

        md: str | None = None
        ext = src.suffix.lower()
        if ext == ".pdf":
            md = pdf_to_markdown(src, title)
        elif ext == ".docx":
            md = docx_to_markdown(src, title)
        elif ext == ".doc":
            md = doc_to_markdown(src, title)
        elif ext == ".xlsx":
            md = xlsx_to_markdown(src, title)
        else:
            print(f"[SKIP] {title}：不支持的格式 {ext}")
            return

        if md is None:
            return

        out_path.write_text(md, encoding="utf-8")
        stats.append((str(out_path.relative_to(ROOT)), title, len(md)))

    # 1. 人才培养方案 PDF -> data/knowledge/培养方案/
    plan_src = RESOURCES / "人才培养方案"
    if plan_src.exists():
        for f in sorted(plan_src.glob("*.pdf")):
            if not f.name.startswith("~$"):
                _convert_one(f, "培养方案")

    # 2. 办事流程 -> data/knowledge/办事流程/
    process_src = RESOURCES / "办事流程"
    if process_src.exists():
        for f in sorted(process_src.iterdir()):
            if f.suffix.lower() in {".pdf", ".docx", ".doc"} and not f.name.startswith("~$"):
                _convert_one(f, "办事流程")

    # 3. 教学安排 -> data/knowledge/教学安排/
    schedule_src = RESOURCES / "教学安排"
    if schedule_src.exists():
        for f in sorted(schedule_src.iterdir()):
            if f.suffix.lower() in {".pdf", ".docx", ".doc", ".xlsx"} and not f.name.startswith("~$"):
                _convert_one(f, "教学安排")

    # 4. 规章制度 -> data/knowledge/制度流程/
    rule_src = RESOURCES / "规章制度"
    if rule_src.exists():
        for f in sorted(rule_src.iterdir()):
            if f.suffix.lower() in {".pdf", ".docx", ".doc"} and not f.name.startswith("~$"):
                _convert_one(f, "制度流程")

    # 5. 顶层散落文件（兼容旧逻辑）-> data/knowledge/制度流程/
    for f in sorted(RESOURCES.iterdir()):
        if f.is_file() and f.suffix.lower() in {".pdf", ".docx", ".doc"} and not f.name.startswith("~$"):
            _convert_one(f, "制度流程")

    print("\n===== 转换完成 =====")
    for path, title, size in stats:
        print(f"[OK] {title} -> {path}（{size} 字符）")
    if skipped:
        print(f"\n[SKIP] 已存在，跳过 {len(skipped)} 个：")
        for t in skipped:
            print(f"  - {t}")
    print(f"\n共转换 {len(stats)} 个，跳过 {len(skipped)} 个")


if __name__ == "__main__":
    main()
